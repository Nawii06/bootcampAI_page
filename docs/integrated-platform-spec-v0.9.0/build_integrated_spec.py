from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path
from typing import Iterable, Sequence
import json
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

OUT = Path('/mnt/data')
DOCX_PATH = OUT / '국립한국교통대학교_부트캠프AI_통합운영플랫폼_세부명세서_v0.9.0.docx'
MD_PATH = OUT / '국립한국교통대학교_부트캠프AI_통합운영플랫폼_세부명세서_v0.9.0.md'
MANIFEST_PATH = OUT / '국립한국교통대학교_부트캠프AI_통합운영플랫폼_세부명세서_v0.9.0_manifest.json'

BLUE = '0D2E8B'
BURGUNDY = '590000'
CHARCOAL = '1F2937'
DARKGRAY = '4C4C45'
MEDGRAY = 'CCCCCC'
LIGHTGRAY = 'EBEBDF'
SOFTWHITE = 'F7F8FA'
WHITE = 'FFFFFF'

FONT_BODY = 'NanumGothic'
FONT_HEAD = 'NanumSquare'

@dataclass
class TableBlock:
    headers: list[str]
    rows: list[list[str]]
    widths: list[float] | None = None
    note: str | None = None

@dataclass
class SectionBlock:
    title: str
    paragraphs: list[str] = field(default_factory=list)
    bullets: list[str] = field(default_factory=list)
    table: TableBlock | None = None
    code: str | None = None

@dataclass
class Chapter:
    number: str
    title: str
    status: str
    sections: list[SectionBlock]


def p(text: str) -> str:
    return re.sub(r'\s+', ' ', text.strip())


def req_rows(prefix: str, items: Sequence[str], priority: str = 'P1', status: str = 'PARTIAL') -> list[list[str]]:
    return [[f'{prefix}-{i:03d}', item, priority, status] for i, item in enumerate(items, 1)]


def status_table(status: str, blockers: str, evidence: str) -> TableBlock:
    return TableBlock(
        ['구분', '판정'],
        [
            ['현재 상태', status],
            ['주요 차단요인', blockers],
            ['완료 증적', evidence],
        ],
        [3.0, 13.5],
    )


def chapter_common_requirements(prefix: str, items: list[str], status='PARTIAL') -> SectionBlock:
    return SectionBlock(
        '요구사항 목록',
        table=TableBlock(['요구사항 ID', '요구사항', '우선순위', '현재 상태'], req_rows(prefix, items, 'P1', status), [3.2, 9.3, 1.8, 2.8])
    )

# ---------------------------------------------------------------------------
# Core chapters
# ---------------------------------------------------------------------------
chapters: list[Chapter] = []

chapters.append(Chapter('1', '문서 개요', 'COMPLETE', [
    SectionBlock('작성 목적', paragraphs=[p('본 문서는 국립한국교통대학교 첨단산업 인재양성 부트캠프사업(AI)의 공개 홈페이지, 학생 포털, 기업 포털 및 관리자 포털을 하나의 통합운영 플랫폼으로 구축·운영하기 위한 업무·기능·데이터·보안·시험·운영 기준을 정의한다.'), p('문서는 사업계획서, 전문기관 지침, 대학 내부정책, 현재 GitHub 저장소의 구현기준을 함께 고려하되, 공식 승인되지 않은 정책이나 외부연계 정보는 확정사항과 구분하여 표시한다.')]),
    SectionBlock('문서 기본정보', table=TableBlock(['항목', '내용'], [
        ['문서명', '국립한국교통대학교 첨단산업 인재양성 부트캠프사업(AI) 통합운영 플랫폼 세부 명세서'],
        ['버전', 'v0.9.0 검토본'],
        ['작성기준일', '2026-07-30'],
        ['대상 저장소', 'Nawii06/bootcampAI_page'],
        ['코드 기준선', 'main / 2346fb2ed634c94f6d7d978dbda1ec8fe22d93ef (재확인 필요)'],
        ['문서 상태', 'DOCUMENT_BASELINE / 공식 승인 전'],
        ['출시 판정', 'NO-GO'],
    ], [4.0, 12.5])),
    SectionBlock('문서 구성', bullets=['본문 제1장부터 제25장까지의 시스템·업무·운영 명세', '부록 A부터 H까지의 용어, 권한, 화면, 추적, API·데이터, 테스트, 위험, 변경·승인 대장', '요구사항 ID, 화면 ID, API ID, 데이터 ID, 테스트 ID, 위험 ID 간 상호 추적']),
    SectionBlock('기준 문서', table=TableBlock(['출처 ID', '문서'], [
        ['SRC-PLAN-001', '국립한국교통대학교 첨단산업 인재양성 부트캠프사업(AI) 수정사업계획서'],
        ['SRC-GUIDE-001', '첨단산업 인재양성 부트캠프 사업관리 운영지침(2026.3.26. 개정)'],
        ['SRC-QA-001', '첨단산업 인재양성 부트캠프사업 가이드(Q&A)'],
        ['SRC-QA-002', '2025 첨단산업 부트캠프 사업설명회 Q&A'],
        ['SRC-BUD-001', '사업비 교육자료_부트캠프(2026)'],
        ['SRC-CODE-001', 'GitHub 저장소 main 기준 소스코드'],
    ], [3.5, 13.0])),
    chapter_common_requirements('DOC', [
        '문서의 목적, 적용범위, 독자, 기준선 및 승인상태를 명확히 표시해야 한다.',
        '확정사항, 제안사항, 외부정보 대기사항 및 실환경 검증사항을 구분해야 한다.',
        '모든 장과 부록은 고유 식별자와 상호참조체계를 사용해야 한다.',
        '코드·문서·정책·운영 기준선을 독립적으로 관리해야 한다.',
        '문서 변경은 부록 H 변경·승인기록에 연결해야 한다.',
    ], 'COMPLETE'),
]))

chapters.append(Chapter('2', '시스템 목적 및 적용범위', 'PARTIAL', [
    SectionBlock('시스템 목적', paragraphs=[p('통합운영 플랫폼은 사업 홍보만을 위한 홈페이지가 아니라 교육과정, 프로그램, 학생 이수, 수혜, 기업협력, 예산, 성과, 콘텐츠와 증빙을 연결하여 사업단의 운영기록을 일관되게 관리하는 업무시스템이다.'), p('대학 ERP, RCMS, 학사시스템 및 전문기관 시스템의 공식 원장을 대체하지 않고, 업무상 필요한 참조정보와 대조상태를 관리하는 보조·통합 플랫폼으로 한정한다.')]),
    SectionBlock('적용범위', table=TableBlock(['영역', '포함 기능', '시스템 경계'], [
        ['공개 홈페이지', '사업·교육과정·모집·기업·성과·공지·자료', '승인된 공개정보만 제공'],
        ['학생 포털', '신청·학습·이수·수혜·포트폴리오', '학생 본인정보 범위'],
        ['기업 포털', '참여신청·확약·프로젝트·평가·채용', '소속 기업과 배정업무 범위'],
        ['관리자 포털', '교육·수혜·기업·예산·성과·CMS·감사', '역할·Scope·직무분리'],
        ['외부연계', 'SSO·학사·ERP·RCMS·전문기관 / 양단산에 ·공식 원천·원장 유지'],
        ['운영', '배포·모니터링·백업·복구', '운영도구와 Runbook 병행'],
    ], [3.1, 7.0, 6.4])),
    SectionBlock('비적용 범위', bullets=['대학 회계원장의 직접 대체', 'RCMS 집행·정산 기능의 직접 대체', '학사 성적·졸업판정의 원천시스템 대체', '은행 계좌정보와 지급수단 저장', '대학 전체 포털 또는 전자결재의 대체']),
    SectionBlock('주요 성과목표 연계', paragraphs=[p('사업계획상 양성인원, 중·고급 이수자, 배출인원, 취·창업, 연계취업, 참여기업 협업, 현장실습·인턴십 및 만족도 지표를 지원한다. 지표별 공식 정의와 연차 목표는 정책 버전으로 관리하고, 내부 계산값과 전문기관 인정값을 분리한다.')]),
    chapter_common_requirements('DOC-SCP', [
        '시스템은 사업 운영 데이터의 통합 조회와 업무흐름을 지원해야 한다.',
        '공식 학사·회계·전문기관 원장과 내부 관리값을 구분해야 한다.',
        '사업연도별 정책과 데이터를 분리해야 한다.',
        '개인정보와 금전정보는 목적에 필요한 최소 범위만 처리해야 한다.',
        '외부연계가 없을 때 승인된 수동 대체절차를 제공해야 한다.',
    ]),
]))

chapters.append(Chapter('3', '사용자 및 권한 정의', 'PARTIAL', [
    SectionBlock('사용자 유형', table=TableBlock(['유형', '설명', '기본 포털'], [
        ['PUBLIC', '비로그인 공개 이용자', '공개 홈페이지'],
        ['STUDENT', '사업 참여 또는 신청 학생', '학생 포털'],
        ['COMPANY_APPLICANT', '기업 참여신청 사용자', '기업 포털'],
        ['COMPANY_MANAGER', '승인 기업의 업무담당 사용자', '기업 포털'],
        ['STAFF', '교육·수혜·기업·예산·성과·콘텐츠 담당자', '관리자 포털'],
        ['REVIEWER', '검토·승인 담당자', '관리자 포털'],
        ['SYSTEM_ADMIN', '기술·계정·설정 관리자', '관리자 포털'],
        ['AUDITOR', '읽기전용 감사·추적 담당자', '관리자 포털'],
    ], [3.6, 8.3, 4.6])),
    SectionBlock('역할 코드', bullets=['EDUCATION_STAFF: 교과·교육과정·프로그램 운영', 'BENEFIT_STAFF: 수혜정책·후보·지급 대조', 'COMPANY_STAFF: 기업신청·확약·활동 검토', 'BUDGET_STAFF: 예산배정·집행·정산 대조', 'PERFORMANCE_STAFF: 지표·목표·실적·제출', 'CONTENT_EDITOR: 콘텐츠 작성·수정', 'REVIEWER: 지정 업무 검토·승인', 'SYSTEM_ADMIN: 기술설정·계정·운영', 'AUDITOR: 감사로그·증적 읽기전용']),
    SectionBlock('권한 판정 순서', code='인증 → 계정 활성 → 역할 → 역할 유효기간 → 사업연도·학과·기업·프로그램 Scope → 리소스 소유권 → 업무상태 → 직무분리'),
    SectionBlock('직무분리 원칙', bullets=['작성자와 최종 승인자를 분리한다.', '산정자와 지급·집행 승인자를 분리한다.', '시스템 관리자는 기술권한을 이유로 사업업무를 자동 승인하지 못한다.', '감사자는 수정권한을 가지지 않는다.', '긴급권한은 기간·사유·승인·회수 이력을 남긴다.']),
    SectionBlock('현재 구조상 주요 공백', paragraphs=[p('현재 역할 기반 Guard와 사용자 역할 테이블은 존재하지만, 사용자-기업의 명시적 Membership, 학과·프로그램 단위 Scope의 일관된 적용, 검토건 배정, 시스템 관리자 우회범위 통제가 부족하다. user_roles의 기본키에 scope_id가 포함되지 않은 구조는 같은 Scope 유형의 복수대상 부여에 제약이 될 수 있다.')]),
    chapter_common_requirements('ROL', [
        '모든 보호 API는 역할과 Scope를 함께 검증해야 한다.',
        '학생은 본인 데이터만 조회·변경할 수 있어야 한다.',
        '기업 사용자는 승인된 소속기업 데이터만 처리할 수 있어야 한다.',
        '검토자는 배정된 건과 승인범위만 처리해야 한다.',
        '역할부여·변경·만료·회수 이력을 보존해야 한다.',
        '시스템 관리자 우회는 최소화하고 감사대상으로 기록해야 한다.',
    ]),
]))

chapters.append(Chapter('4', '전체 메뉴구조', 'PARTIAL', [
    SectionBlock('정보구조 원칙', bullets=['공개·학생·기업·관리자 포털의 목적과 탐색구조를 분리한다.', '메뉴 노출과 직접 URL 접근권한을 동일한 정책원천에서 관리한다.', '상위 메뉴와 하위 상세경로의 활성상태가 중복되지 않도록 Canonical Route를 지정한다.', '모바일·접근성·Breadcrumb를 공통 적용한다.']),
    SectionBlock('최상위 메뉴', table=TableBlock(['포털', '최상위 메뉴'], [
        ['공개', '홈 / 사업소개 / 교육과정 / 모집·참여 / 참여기업 / 성과·소식 / 자료실'],
        ['학생', '대시보드 / 프로그램 / 학습활동 / 이수현황 / 수혜현황 / 포트폴리오 / 내 정보'],
        ['기업', '대시보드 / 참여신청 / 기업정보 / 확약·활동 / 프로젝트 / 평가·채용 / 알림'],
        ['관리자', '대시보드 / 교육 / 수혜 / 기업 / 예산 / 성과 / CMS / 파일 / 감사 / 시스템'],
    ], [3.0, 13.5])),
    SectionBlock('라우팅 기준', bullets=['공개: /, /about, /curriculum, /programs, /companies, /performance, /content', '학생: /student/*', '기업: /partner/*', '관리자: /admin/*', '시스템 상태: /api/healthz, /api/readyz, /api/metrics', '업무 API: /api/v1/*']),
    SectionBlock('관리원칙', paragraphs=[p('메뉴는 역할별로 숨기더라도 API는 독립적으로 권한을 재검증해야 한다. 상세화면과 감사로그 Deep Link는 동일한 리소스 식별자를 사용하고, 폐기 라우트는 대체경로와 종료일을 제공해야 한다.',  'PUB'),
    domain_chapter(6, '학생 포털 화면 명세', 'PARTIAL',
        '학생이 프로그램 신청부터 학습, 이수, 수혜 및 포트폴리오까지 본인의 참여현황을 확인하고 필요한 조치를 수행할 수 있도록 한다.',
        ['STUDENT', '교육담당자(문의·정정 처리)', '수혜담당자(수혜상태 처리)'],
        ['대시보드', '프로그램 조회·신청·취소·보완', '출석·과제·설문', '이수판정·부족요건', '수혜후보·승인·지급상태', '경험기록·포트폴리오 공유', '동의·알림·내 정보'],
        ['학생 프로필', '프로그램 신청', '출석·과제·설문', '교과·비교과 이수', '이수판정 Snapshot', '수혜후보·지급상태', '경험기록·공유토큰'],
        ['모든 학생 데이터는 로그인 계정과 연결된 학생 ID로 서버에서 확정한다.', '타 학생 ID를 Query·Path로 전달하더라도 접근을 차단한다.', '학사 원천필드는 학생이 직접 수정하지 못한다.', '공유토큰은 만료·철회·Rate Limit을 적용한다.'],
        '기본 화면과 일부 학생 소유권 검사는 구현되어 있으나 이수·수혜 이의신청, 동의관리, 알림, 학사 원천 프로필 관리와 실제 SSO 연결이 미완료이다.', 'STU'),
    domain_chapter(7, '기업 포털 화면 명세', 'PARTIAL',
        '참여기업이 기업신청, 정보관리, 확약, 프로젝트, 전문가, 학생평가 및 채용연계 실적을 안전하게 처리하도록 한다.',
        ['COMPANY_APPLICANT', 'COMPANY_MANAGER', 'COMPANY_STAFF', 'REVIEWER'],
        ['기업 참여신청·보완', '기업정보·담당자·전문가', '참여확약', '수요조사·프로젝트 제안', '학생 배정·평가', '취업·활동실적', '알림'],
        ['기업신청', '기업마스터', '기업 담당자·전문가', '연도별 확약', '기업 활동·연결학생', '고용연계 결과', '기업 파일'],
        ['신청자는 본인 신청만 조회한다.', '승인 후 기업마스터를 생성하고 신청과 연결한다.', '기업 사용자는 명시적인 Membership과 유효기간을 가져야 한다.', '기업에 학생정보를 제공할 때 배정관계와 학생동의를 검증한다.'],
        '기업신청·기업마스터·담당자·전문가·확약·활동 API는 존재하지만 사용자-기업 Membership, 복수 사용자 초대·회수, 수요조사·평가·채용의 완결된 Workflow가 부족하다.', 'PRT'),
    domain_chapter(8, '관리자 포털 화면 명세', 'PARTIAL',
        '역할별 담당자가 교육, 수혜, 기업, 예산, 성과, 콘텐츠, 파일, 감사 및 시스템 업무를 통제된 절차로 처리하도록 한다.',
        ['EDUCATION_STAFF', 'BENEFIT_STAFF', 'COMPANY_STAFF', 'BUDGET_STAFF', 'PERFORMANCE_STAFF', 'CONTENT_EDITOR', 'REVIEWER', 'SYSTEM_ADMIN', 'AUDITOR'],
        ['역할별 대시보드와 할 일', '도메인별 목록·상세·등록·수정', '검토·승인·반려', '사업연도·학과·기업 Scope', '대량처리·내보내기', '변경이력·감사 Deep Link'],
        ['도메인 업무데이터', '역할·Scope', '검토배정', '감사로그', '내보내기 기록', '운영·시스템 설정'],
        ['메뉴 노출과 직접 URL 권한을 일치시킨다.', '작성자·검토자·승인자 직무를 분리한다.', '대량 내보내기는 목적·승인된 금한다.', '감사자는 읽기전용으로 운영한다.'],
        '도메인별 주요 관리자 화면이 존재하나고 통합하는 일, 검토배정, 학과·기업 Scope, 권한관리, 운영·위험·변경 관리화면은 목표상태이다.', 'ADM'),
    domain_chapter(9, '교과목·교육과정 기능 명세', 'PARTIAL',
        '교과목 마스터, 개설교과목, 교육과정 버전과 이수요건을 관리하고 학사 원천자료를 안전하게 반영한다.',
        ['EDUCATION_STAFF', 'REVIEWER', 'AUDITOR'],
        ['교과목·개설교과 CRUD', '교육과정 버전·유효기간', '이수요건 등록', 'CSV·XLSX·JSON·외부 API Import', 'Preview·Commit·오류행', '대체교과·정정'],
        ['course_masters', 'course_offerings', 'curricula', 'curriculum_requirements', 'import_jobs', 'import_staging_rows'],
        ['동일 파일과 외부키의 중복반영을 차단한다.', 'Import는 Stage-Preview-Commit 단계로 처리한다.', '게시·확정된 과거 교육과정 버전을 보존한다.', '학사 정정 시 영향대상과 재계산을 승인한다.'],
        '기본 교과·교육과정과 Import API는 구현되어 있으나 실제 학사연계, 대체·선수·복합 선택군의 정규화, 성적정정 수신과 데이터품질 대장이 미완료이다.', 'CUR'),
    domain_chapter(10, '비교과 프로그램 기능 명세', 'PARTIAL',
        '비교과·몰입형·산학연계 프로그램의 기획, 모집, 신청, 선발, 운영, 출석, 과제, 설문 및 이수확정을 관리한다.',
        ['STUDENT', 'EDUCATION_STAFF', 'REVIEWER', '기업 전문가'],
        ['프로그램·회차 등록', '모집기간·정원·자격', '신청·대기·선발', '출석·과제·설문', '이수확정', '운영결과·성과연계'],
        ['programs', 'program_sessions', 'program_applications', 'attendance_events', 'attendance_records', 'assignments', 'assignment_submissions', 'surveys', 'survey_responses', 'program_completions'],
        ['모집·운영기간과 상태전이를 검증한다.', '학생-회차 중복신청과 동시 정원초과를 DB Transaction으로 차단한다.', '출석·과제·설문 정정은 사유·승인 이력을 보존한다.', '이수확정 시 규칙과 원천값 Snapshot을 저장한다.'],
        '프로그램 운영 전반의 API와 데이터는 존재하지만 브라우저 E2E, 정원 동시성 실 DB 시험, 대기자·취소·보완, 운영결과 보고가 미완료이다.', 'PRG'),
    domain_chapter(11, '학생 이수판정 기능 명세', 'PARTIAL',
        '교과·비교과·프로젝트·현장실습·인턴십 원천과 교육과정 규칙을 결합하여 학생별 이수상태와 부족요건을 산출한다.',
        ['STUDENT', 'EDUCATION_STAFF', 'REVIEWER', 'AUDITOR'],
        ['규칙 버전', 'DB 원천 기반 판정', '충족·미충족 요건', '진행률·자료대기', '판정 Snapshot', '확정·정정·재계산·이의신청'],
        ['course_completions', 'program_completions', 'experiential_records', 'completion_assessments', 'curriculum_requirements'],
        ['양성인원은 해당연차 1개 이상 교육프로그램 이수자를 학생 단위로 중복 제거한다.', '중·고급 이수자 비율은 승인된 수준코드와 프로그램 이수확정을 기준으로 한다.', '확정결과는 정책변경만으로 자동 소급하지 않는다.', '자료 미수신·미확정과 미충족을 구분한다.'],
        '판정 계산과 Snapshot 구조는 존재하지만 공식 이수정책, 확정상태·승인자·현재 유효본, 정정·이의신청과 실제 학사성적 검증이 미완료이다.', 'CMP'),
    domain_chapter(12, '장학금·수혜 기능 명세', 'PARTIAL',
        '사업연도별 수혜정책과 자격조건을 버전화하고 학생별 후보산정, 검토, 승인, 지급대조 및 환수를 관리한다.',
        ['BENEFIT_STAFF', 'REVIEWER', 'STUDENT', 'AUDITOR'],
        ['수혜정책·조건', '개별·일괄 후보산정', 'Dry-run', '승인·반려', '지급상태·ERP 참조', '중복수혜·취소·환수', '학생 수혜조회'],
        ['benefit_policies', 'benefit_eligibility_rules', 'benefit_candidates', 'benefit_approvals', 'benefit_payments', '향후 benefit_recoveries'],
        ['금액·산식은 코드에 고정하지 않고 정책버전으로 관리한다.', '후보산정은 원천값·정책·결과 Snapshot을 보존한다.', '산정자와 최종 승인자를 분리한다.', '실제 지급은 ERP를 원장으로 하고 플랫폼은 요청·상태·대조를 관리한다.', '계좌정보는 플랫폼에 저장하지 않는다.'],
        '정책·후보·승인·지급 API와 5개 테이블은 존재하지만 대학 공식 정책, 외부 중복수혜, ERP 지급연계, 지급취소·환수, 학생 이의신청과 실환경 동시성 검증이 미완료이다.', 'BEN'),
    domain_chapter(13, '참여기업 기능 명세', 'PARTIAL',
        '기업 참여신청부터 승인기업 마스터, 담당자·전문가, 확약, 활동·프로젝트·학생평가·채용실적을 연계한다.',
        ['COMPANY_APPLICANT', 'COMPANY_MANAGER', 'COMPANY_STAFF', 'REVIEWER', 'AUDITOR'],
        ['참여신청·보완·결정', '기업마스터·중복관리', '담당자·전문가', '연도별 확약', '활동·프로젝트·학생 연결', '공개기업·성과연계'],
        ['company_applications', 'companies', 'company_contacts', 'company_experts', 'company_commitments', 'company_participations', '향후 company_memberships'],
        ['사업자등록번호와 외부키로 중복을 통제한다.', '승인 신청과 기업마스터의 관계를 보존한다.', '대표 담당자·사용자 소속·유효기간을 관리한다.', '학생정보 제공은 동의·배정관계·목적을 검증한다.'],
        '기업 6개 테이블과 20개 Endpoint는 존재하지만 사용자-기업 Membership, 기업 초대·회수, 중복기업 병합, 세분화된 프로젝트·평가·고용 모델이 미완료이다.', 'COM'),
    domain_chapter(14, '예산관리 기능 명세', 'PARTIAL',
        '사업연도·프로그램별 예산배정, 집행, 변경, 증빙과 ERP·RCMS 참조·대조상태를 관리하는 보조장부를 제공한다.',
        ['BUDGET_STAFF', 'REVIEWER', 'AUDITOR'],
        ['예산계획·배정', '집행등록·증빙', '금액변경·이력', '잔액·집행률', 'ERP·RCMS 참조·대조', '취소·불인정·환수·이월·정산'],
        ['budget_allocations', 'budget_executions', 'budget_change_history', '향후 budget_plans·reconciliations·settlements'],
        ['플랫폼은 공식 회계원장을 대체하지 않는다.', '집행 Transaction 안에서 잔액을 재검증하고 동시 초과를 차단한다.', '모든 금액변경은 전후값·사유·승인을 기록한다.', 'ERP·RCMS 참조번호는 실제 원장과 대조되어야 한다.'],
        '배정·집행·변경 API와 데이터는 존재하지만 실제 PostgreSQL 동시성 시험, ERP·RCMS 연계·대조, 이월·정산·불인정·환수 Workflow가 미완료이다.', 'BGT'),
    domain_chapter(15, '성과·증빙 기능 명세', 'PARTIAL',
        '사업계획과 전문기관 기준에 따른 지표, 연차 목표, 원천자료, 실적계산, 증빙, 검토, 공개 및 제출·인정값을 관리한다.',
        ['PERFORMANCE_STAFF', 'REVIEWER', 'AUDITOR', 'PUBLIC'],
        ['지표·산식·목표 버전', '원천자료 요약', 'COUNT·RATE 계산', '실적 Snapshot', '증빙연결', '검토·공개승인', '전문기관 제출·인정값'],
        ['performance_indicators', 'performance_targets', 'performance_results', 'performance_evidence', 'performance_reviews', '향후 submissions·accepted_results'],
        ['지표별 분모·분자·중복제거·기준일을 명시한다.', '내부 계산값, 제출값, 전문기관 인정값을 분리한다.', '확정 실적은 원천·산식·증빙 Snapshot을 보존한다.', '공개 승인된 확정값만 홈페이지에 노출한다.'],
        '지표·목표·결과·증빙·검토 API는 존재하지만 공식 산식 최종 승인, 제출버전·정정·인정값, 질적지표 다중평가와 전문기관 Interface가 미완료이다.', 'PER'),
    domain_chapter(16, 'CMS 기능 명세', 'PARTIAL',
        '공지, 뉴스, 모집, 자료, 성과사례 등 공개 콘텐츠를 작성·검토·승인·게시·예약·보관하고 버전을 관리한다.',
        ['CONTENT_EDITOR', 'REVIEWER', 'AUDITOR', 'PUBLIC'],
        ['콘텐츠 작성·수정', 'slug·메타데이터', '검토·승인·게시·보관', '예약게시', '첨부파일', '버전·복원', '접근성·개인정보 검토'],
        ['content_items', 'content_attachments', 'content_versions', 'stored_files'],
        ['작성자는 자기 게시를 할 수 없다.', '게시 전 개인정보·저작권·접근성 검토를 적용한다.', '예약게시 실패를 탐지·재처리한다.', '게시취소 시 캐시·검색·공개파일 접근을 함께 회수한다.'],
        '콘텐츠 3개 테이블과 6개 Endpoint는 존재하지만 공개 단일 상세, 예약게시 Worker, 개인정보·접근성 검토 Gate, 공개 첨부파일 모델이 미완료이다.', 'CMS'),
    domain_chapter(17, '파일관리 기능 명세', 'PARTIAL',
        '교육·기업·예산·성과·CMS 등에서 생성되는 파일을 안전하게 업로드, 검사, 저장, 연결, 다운로드, 보존, 법적보존, 파기 및 복구한다.',
        ['업무담당자', '학생', '기업 사용자', 'SYSTEM_ADMIN', 'AUDITOR'],
        ['확장자·MIME·Signature 검증', '크기·SHA-256', '악성코드 검사·격리', '로컬·S3 Adapter', '업무관계 권한', '다운로드 감사', '보존·법적보존·파기·복구'],
        ['file_retention_policies', 'stored_files', '향후 file_relationships·scan_results·purge_jobs'],
        ['운영환경은 악성코드 검사 실패 시 Fail Closed를 적용한다.', '파일 목록은 업무관계와 Scope로 제한한다.', '공개여부는 콘텐츠 승인과 분리한다.', 'DB 메타데이터와 Object의 파기·복구 일치를 검증한다.'],
        '업로드·다운로드·보존·법적보존·정리 API는 존재하지만 실제 S3·검사서비스·격리, 범용 파일관계, 공개파일, 전체 복구 증적이 미완료이다.', 'FIL'),
    domain_chapter(18, '인증·세션·권한 명세', 'PARTIAL / BLOCKED_EXTERNAL',
        '대학 통합인증과 서버측 세션을 기반으로 사용자를 식별하고 역할·Scope·소유권을 모든 보호 API에 일관되게 적용한다.',
        ['모든 로그인 사용자', 'SYSTEM_ADMIN', 'AUDITOR'],
        ['SSO 시작·Callback', 'Issuer·Subject 사용자 매핑', '세션 생성·연장·만료·로그아웃·회수', 'Cookie·CSRF', '역할·Scope·유효기간', '고위험 재인증', '인증감사'],
        ['users', 'roles', 'user_roles', 'students', '향후 sessions·role_assignment_history·access_requests'],
        ['운영환경 Mock 인증을 차단한다.', '이메일이 아닌 불변 외부식별자를 사용한다.', 'Secure·HttpOnly·SameSite Cookie와 CSRF 통제를 적용한다.', '시스템 관리자 우회범위를 최소화한다.'],
        '비운영 Mock 인증과 Role Guard는 존재하지만 실제 대학 SSO, 중앙 세션·로그아웃, Claim Mapping, MFA·재인증, Scope 전면적용은 외부정보 대기 또는 미구현이다.', 'AUT'),
    domain_chapter(19, '데이터 모델 및 관계 명세', 'PARTIAL / REAL_ENV_REQUIRED',
        '통합운영 플랫폼의 53개 현행 물리테이블과 목표 엔터티의 구조, 관계, 제약, 버전, Snapshot, 삭제·파기 및 원천정보를 정의한다.',
        ['개발·DBA·업무설계자', '감사·운영 담당자'],
        ['UUID PK', '사업연도 Scope', 'FK·Unique·Check', 'Soft Delete', 'Snapshot·Version', '외부키·원천시스템', 'JSONB Schema Version', 'Migration·Seed'],
        ['공통 4', '사용자·권한 4', '교과 4', 'Import 2', '감사 1', '파일 2', '프로그램 10', '이수 3', '수혜 5', '기업 6', '예산 3', '성과 5', 'CMS 3', 'Rate Limit 1'],
        ['중요 제약은 애플리케이션 검증만이 아니라 DB에 적용한다.', '확정 산정결과는 입력과 규칙 Snapshot을 보존한다.', 'Soft Delete와 실제 개인정보 파기를 구분한다.', '다형 관계와 UUID 배열은 관계테이블로 보완한다.'],
        '현행 Drizzle 스키마는 53개 테이블을 정의하지만 실제 PostgreSQL Migration·Seed·제약·동시성·백업복구 검증이 필요하며 기업 Membership, 동의·이의·대조·제출·운영 엔터티가 추가되어야 한다.', 'DAT'),
    domain_chapter(20, 'API 및 외부연계 명세', 'PARTIAL / BLOCKED_EXTERNAL',
        '현재 116개 API Endpoint와 목표 외부연계를 공통 계약, 권한, 멱등성, 재시도, 감사 및 대조기준으로 관리한다.',
        ['포털 클라이언트', '외부시스템', '운영·개발 담당자'],
        ['/api/v1 버전경로', 'Zod 계약·공통 오류', 'Request ID', '페이지·필터', '멱등성·낙관적 Lock', '비동기 Job·Outbox·Inbox', 'SSO·학사·ERP·RCMS·전문기관·알림'],
        ['현재 116개 Endpoint', 'api-zod 계약', 'integration_operations·outbox·inbox 목표 데이터', '외부 요청·응답 Metadata'],
        ['외부호출은 허용 Host·HTTPS·Timeout·제한 재시도를 적용한다.', '금전·신청·제출은 Idempotency-Key를 적용한다.', '내부상태와 외부상태를 정기 대조한다.', '전체 Endpoint를 OpenAPI에 반영한다.'],
        '현재 Router·Zod·오류응답·Request ID·일부 역할검증은 존재하지만 전체 OpenAPI, 멱등성, Outbox·Inbox, Job, 실제 SSO·학사·ERP·RCMS·전문기관 Adapter는 미완료이다.', 'API'),
    domain_chapter(21, '개인정보·보안·감사 명세', 'PARTIAL / BLOCKED_EXTERNAL',
        '개인정보의 적법한 최소처리, 역할·필드 통제, 안전한 저장·전송, 보존·파기 및 변경·다운로드 감사체계를 정의한다.',
        ['개인정보 담당자', '정보보안 담당자', 'SYSTEM_ADMIN', 'AUDITOR', '전 업무담당자'],
        ['처리목록·동의 버전', '최소수집·마스킹', '정보주체 권리요청', '제3자 제공·위탁', 'TLS·Secret·Rate Limit', '보안이벤트·사고대응', '감사로그·내보내기'],
        ['users·students·company_contacts 등 PII', 'audit_logs', 'stored_files', '향후 consent_records·privacy_requests·disclosures·security_events'],
        ['감사 전후값에 민감정보를 중복저장하지 않는다.', '대량 다운로드는 목적·승인·암호화·만료를 적용한다.', 'Secret 원문은 로그·변경기록에 저장하지 않는다.', '개인정보 사고는 탐지·격리·보고·통지·재발방지를 기록한다.'],
        '감사로그·보안헤더·일부 Rate Limit과 파일 통제는 존재하지만 공식 처리목록·동의·보존기간, 권리요청, 제공·위탁, 중앙 로그·변조방지, 사고대응은 미완료이다.', 'PRV'),
    domain_chapter(22, '오류·예외처리 명세', 'PARTIAL',
        '사용자 입력, 권한, 상태충돌, 외부연계, 부분실패, 비동기 처리 및 업무예외를 일관된 코드와 복구절차로 처리한다.',
        ['모든 사용자', '업무담당자', '운영자', '개발자'],
        ['공통 오류응답', '필드 오류', '401·403·404·409', '외부 Timeout·재시도', '부분실패·배치결과', '예외승인·이의신청', '보상처리·Dead Letter'],
        ['error catalog', 'request_id', 'job items', 'integration operations', 'appeals·exceptions 목표 데이터'],
        ['오류메시지는 민감한 내부정보를 노출하지 않는다.', '재시도 가능·불가능 오류를 구분한다.', '외부 성공 후 내부 실패는 보상·대조 대상으로 등록한다.', '사용자 이의신청과 기술예외를 구분한다.'],
        'Zod·ApiError 기반 공통 형식은 존재하지만 중앙 오류코드 Registry, 비동기 작업·부분실패, 예외승인·이의·보상·Dead Letter 모델이 미완료이다.', 'ERR'),
    domain_chapter(23, '테스트 및 인수기준', 'PARTIAL / REAL_ENV_REQUIRED',
        '요구사항별 단위·계약·통합·E2E·보안·성능·접근성·UAT·OAT를 계획하고 출시 Gate를 정의한다.',
        ['개발자', 'QA', '업무 인수자', '개인정보·보안 담당자', '운영 인수자'],
        ['정적·단위·계약', '실 PostgreSQL 통합·Migration', '브라우저 E2E', '권한·IDOR·보안', '성능·접근성·호환성', 'UAT·OAT·Go/No-Go'],
        ['test cases·runs·results 목표 데이터', 'CI 로그', '결함·증적', 'UAT·OAT 승인서'],
        ['P0 요구사항은 시험설계·실행·통과 100%를 요구한다.', 'Critical·Major 결함은 0건이어야 한다.', '미실행·외부대기 시험을 통과로 간주하지 않는다.', '코드 구현자와 최종 인수자를 분리한다.'],
        '루트 verify는 일반시험·Build·Secret·Container·SBOM를 포함하지만 API integration 시험은 별도이며 CI의 PostgreSQL Service·브라우저 E2E·UAT·OAT가 부족하다.', 'TST'),
    domain_chapter(24, '배포·백업·복구 명세', 'PARTIAL / REAL_ENV_REQUIRED',
        '불변 릴리스, 안전한 Migration, 상태점검, 모니터링, DB·파일 백업, 전체 복구, RPO·RTO 및 재해대응을 정의한다.',
        ['개발책임자', '운영담당자', 'DBA', '정보보안 담당자', '변경승인자'],
        ['Release Manifest·Image Digest', 'SBOM·취약점', '배포 준비검사', 'Migration 계정분리', 'Health·Readiness·Graceful Shutdown', 'DB·파일·Secret Backup', 'Restore·Rollback·DR'],
        ['releases·deployments·backups·restore_rehearsals 목표 데이터', '운영 로그·Metrics', '백업 파일·해시·Manifest'],
        ['운영 Secret은 외부 관리하고 Image에 포함하지 않는다.', 'DB Migration과 Runtime 계정을 분리한다.', '백업은 무결성을 검증하고 분리된 위치에 보관한다.', '빈 환경 전체 복구로 RPO·RTO를 실측한다.'],
        'Container·Compose·준비검사·Backup/Restore 스크립트와 Release 문서는 존재하지만 실제 Registry 배포, 자동 백업, PITR, S3 복제, 중앙 모니터링, 전체 복구 리허설과 DR이 미완료이다.', 'OPS'),
    domain_chapter(25, '미구현·외부정보 대기사항', 'BLOCKED_EXTERNAL / REAL_ENV_REQUIRED',
        '운영 차단요인과 미구현·부분구현·외부정보 대기·실환경 검증 과제를 단일 Gap Register로 관리한다.',
        ['사업총괄', '도메인 업무담당자', '개발·QA·운영', '대학 관련부서', '전문기관·공급자'],
        ['P0·P1·P2 분류', '외부정보 요청·회신기한', '책임자·선행조건·완료증적', '의존성·위험·변경 연결', 'Release Gate·Go/No-Go'],
        ['gap_items·external_dependencies 목표 데이터', '요청공문·회신', 'Issue·PR·시험증적', '인수·위험수용 기록'],
        ['P0는 운영 전 0건이어야 한다.', '외부정보 수신만으로 종료하지 않고 구현·통합시험을 완료한다.', 'P1 조건부 수용은 책임자·기한·대체절차·모니터링을 요구한다.', 'P0 보안·금전·복구 위험은 예외승인으로 우회하지 않는다.'],
        '대학 SSO, 학사·ERP·RCMS·전문기관 Interface, 역할 Scope, 기업 Membership, 개인정보 정책, 실제 PostgreSQL·S3·검사서비스, E2E·UAT·OAT·복구훈련이 현재 핵심 P0이다.', 'GAP'),
])

# ---------------------------------------------------------------------------
# Appendices
# ---------------------------------------------------------------------------
appendices: list[Chapter] = []

# A Glossary
terms = [
    ['양성인원', '해당연차 1개 이상 교육프로그램을 이수 완료하여 사업단 명의 이수증 또는 디지털 배지를 받은 학생 수. 동일연차 중복 불인정.'],
    ['중·고급 이수자', '양성인원 중 승인된 중급 또는 고급 프로그램을 이수 완료한 학생. 공식 수준코드와 이수확정 자료로 산정.'],
    ['배출인원', '양성인원 중 해당연차 졸업(예정)자. 전문기관 정의와 기준일에 따라 산정.'],
    ['교육프로그램', '교과형, 비교과형, 몰입형, 프로젝트, 현장실습 등 사업단이 이수요건을 정하여 운영하는 교육단위.'],
    ['교과목', '대학 학사체계에 등록된 정규 교과목 마스터 및 학기별 개설정보.'],
    ['참여기업', '주관대학과 협약 또는 참여확약을 통해 교육과정 개발·운영, 프로젝트, 현장실습, 채용연계 등에 참여하는 기업.'],
    ['K-PASS', '전문기관의 사업 공고·협약·과제·성과 등 전주기를 관리하는 사업관리시스템.'],
    ['RCMS', '사업비 지급·사용·관리·정산에 활용되는 실시간통합연구비관리시스템.'],
    ['ERP', '대학의 회계·지급·전표 등 공식 원장 시스템. 본 플랫폼은 ERP를 대체하지 않음.'],
    ['Snapshot', '특정 시점의 입력값, 규칙, 계산결과를 재현할 수 있도록 고정 저장한 데이터.'],
    ['Scope', '역할이 효력을 가지는 사업연도·학과·기업·프로그램·검토건 등의 범위.'],
    ['IDOR', '식별자 변경으로 권한 없는 다른 사용자·기업의 데이터에 접근하는 취약점.'],
    ['Soft Delete', '삭제일을 기록하여 일반 조회에서 제외하는 논리삭제. 실제 개인정보 파기와 구분.'],
    ['P0/P1/P2', 'P0 운영차단, P1 개톴 전 핵심완성, P2 개선·최적화 우선순위.'],
    ['UAT', '업무담당자가 실제 업무정책과 절차 충족 여부를 확인하는 사용자 인수시험.'],
    ['OAT', '운영담당자가 배포·감시·백업·복구 가능성을 확인하는 운영 인수시험.'],
    ['RPO', '장애 시 허용 가능한 데이터 손실 시점 목표.'],
    ['RTO', '장애 후 서비스 복구까지 허용 가능한 시간 목표.'],
    ['Outbox/Inbox', 'DB Transaction과 외부메시지 전송·수신의 유실·중복을 통제하는 통합 패턴.'],
    ['Go/No-Go', '정식운영 여부를 P0, 결함, 인수, 복구, 위험을 기준으로 판단하는 최종 결정.'],
]
appendices.append(Chapter('A', '용어·약어집', 'COMPLETE', [
    SectionBlock('사업·업무 용어', table=TableBlock(['용어', '정의'], terms, [4.2, 12.3])),
    SectionBlock('주요 약어', table=TableBlock(['약어', '의미'], [
        ['API', 'Application Programming Interface'], ['CMS', 'Content Management System'], ['SSO', 'Single Sign-On'], ['OIDC', 'OpenID Connect'], ['PII', 'Personally Identifiable Information'], ['RBAC', 'Role-Based Access Control'], ['CI/CD', 'Continuous Integration / Continuous Delivery'], ['SBOM', 'Software Bill of Materials'], ['PITR', 'Point-in-Time Recovery'], ['SLA', 'Service Level Agreement'], ['ADR', 'Architecture Decision Record'], ['DR', 'Disaster Recovery'],
    ], [4.2, 12.3])),
    SectionBlock('구현상태 코드', table=TableBlock(['상태', '정의'], [
        ['COMPLETE', '문서·설계·구현·검증이 해당 범위에서 완료'], ['PARTIAL', '일부 구현 또는 증적 부족'], ['PREVIEW_ONLY', '시연·가상데이터 용도'], ['CODE_COMPLETE', '코드는 있으나 실환경 검증 전'], ['REAL_ENV_REQUIRED', '실제 DB·SSO·스토리지·외부시스템 검증 필요'], ['BLOCKED_EXTERNAL', '외부기관 정보·환경·승인 대기'], ['NOT_IMPLEMENTED', '미구현'], ['DEFERRED', '승인된 보류'], ['VERIFIED_OPERATIONAL', '운영환경 적용·검증 완료'],
    ], [4.2, 12.3])),
]))

# B Roles matrix
roles = ['STUDENT','COMPANY_APPLICANT','COMPANY_MANAGER','EDUCATION_STAFF','BENEFIT_STAFF','COMPANY_STAFF','BUDGET_STAFF','PERFORMANCE_STAFF','CONTENT_EDITOR','REVIEWER','SYSTEM_ADMIN','AUDITOR']
appendices.append(Chapter('B', '역할·권한 매트릭스', 'PARTIAL', [
    SectionBlock('역할 정의', table=TableBlock(['역할', '주요 책임', '제한'], [
        ['STUDENT', '본인 신청·학습·이수·수혜·포트폴리오', '타 학생·관리자 업무 금지'],
        ['COMPANY_APPLICANT', '본인 기업 참여신청·보완', '승인기업 업무 금지'],
        ['COMPANY_MANAGER', '소속기업 정보·확약·활동·평가', '타 기업 접근 금지'],
        ['EDUCATION_STAFF', '교과·교육과정·프로그램·이수', '수혜·예산 최종승인 금지'],
        ['BENEFIT_STAFF', '수혜정책·후보·지급상태', '산정자 자기승인 금지'],
        ['COMPANY_STAFF', '기업신청·기업마스터·활동 검토', '소속기업 사용자 역할과 구분'],
        ['BUDGET_STAFF', '배정·집행·대조', '공식 원장 임의수정 금지'],
        ['PERFORMANCE_STAFF', '지표·목표·실적·증빙', '자기 공개승인 금지'],
        ['CONTENT_EDITOR', '콘텐츠 작성·수정', '자기 게시 금지'],
        ['REVIEWER', '배정된 도메인 검토·승인', '전역·자기승인 금지'],
        ['SYSTEM_ADMIN', '계정·설정·운영·기술관리', '사업업무 자동승인 금지'],
        ['AUDITOR', '감사·증적 읽기전용', '수정·승인 금지'],
    ], [3.6, 8.0, 4.9])),
    SectionBlock('도메인 권한 매트릭스', table=TableBlock(['도메인', '조회', '작성·수정', '검토·승인', '감사'], [
        ['교과·프로그램', '학생 본인/교육/검토/감사', '교육담당', '지정 검토자', 'AUDITOR'],
        ['이수판정', '학생 본인/교육/검토/감사', '교육담당 계산·정정', '지정 검토자', 'AUDITOR'],
        ['장학금·수혜', '학생 본인/수혜/검토/감사', '수혜담당', '분리된 검토자', 'AUDITOR'],
        ['기업', '소속기업/기업담당/검토/감사', '소속기업·기업담당', '지정 검토자', 'AUDITOR'],
        ['예산', '예산/검토/감사', '예산담당', '결재권자·검토자', 'AUDITOR'],
        ['성과', '성과/검토/감사/공개', '성과담당', '지정 검토자', 'AUDITOR'],
        ['CMS', '편집/검토/감사/공개', '편집자', '검토자', 'AUDITOR'],
        ['파일', '업무관계·본인', '업로더·업무담당', '공개·파기 승인', 'AUDITOR'],
        ['시스템', 'SYSTEM_ADMIN/AUDITOR', 'SYSTEM_ADMIN', '운영승인자', 'AUDITOR'],
    ], [3.0, 4.5, 3.2, 3.2, 2.6])),
    SectionBlock('Scope 유형', bullets=['GLOBAL: 제한적으로만 사용', 'BUSINESS_YEAR: 사업연도', 'DEPARTMENT: 학과·부서', 'COMPANY: 기업', 'PROGRAM: 프로그램·회차', 'REVIEW_ASSIGNMENT: 검토건', 'RESOURCE_OWNER: 학생 본인·업로더']),
    SectionBlock('확정 필요사항', bullets=['사용자-기업 Membership 모델', '다중 Scope 역할배정 PK', '검토배정과 승인단계', '시스템 관리자 우회범위', '긴급권한 승인·회수', '대량 내보내기 승인권자']),
]))

# C routes - target list
route_rows = [
    ['공개','SCR-PUB-HOM-001','홈','/','현재/보완'],
    ['공개','SCR-PUB-INT-001','사업소개','/about','현재/보완'],
    ['공개','SCR-PUB-CUR-001','교육과정','/curriculum','현재/보완'],
    ['공개','SCR-PUB-REC-001','모집·프로그램','/programs','현재/보완'],
    ['공개','SCR-PUB-COM-001','참여기업','/companies','현재/보완'],
    ['공개','SCR-PUB-PER-001','성과','/performance','현재/보완'],
    ['공개','SCR-PUB-CMS-001','공지·소식','/content','목표'],
    ['공개','SCR-PUB-RES-001','자료실','/resources','현재/보완'],
    ['공개','SCR-PUB-POR-001','공유 포트폴리오','/portfolio/:token','현재/보완'],
    ['학생','SCR-STU-HOM-001','대시보드','/student','현재'],
    ['학생','SCR-STU-APP-001','프로그램 신청','/student/programs','현재/보완'],
    ['학생','SCR-STU-LRN-001','학습활동','/student/learning','현재/보완'],
    ['학생','SCR-STU-CMP-001','이수현황','/student/completion','현재/보완'],
    ['학생','SCR-STU-BEN-001','수혜현황','/student/benefits','목표/부분'],
    ['학생','SCR-STU-POR-001','포트폴리오','/student/portfolio','현재/보완'],
    ['학생','SCR-STU-PRF-001','내 정보·동의','/student/profile','목표'],
    ['기업','SCR-PRT-HOM-001','대시보드','/partner','현재'],
    ['기업','SCR-PRT-APP-001','참여신청','/partner/application','현재/보완'],
    ['기업','SCR-PRT-COM-001','기업정보','/partner/company','현재/보완'],
    ['기업','SCR-PRT-SRV-001','수요조사','/partner/surveys','목표'],
    ['기업','SCR-PRT-PRJ-001','프로젝트·활동','/partner/projects','현재/보완'],
    ['기업','SCR-PRT-EVL-001','학생평가','/partner/evaluations','목표/부분'],
    ['기업','SCR-PRT-EMP-001','채용연계','/partner/employment','목표/부분'],
    ['관리자','SCR-ADM-HOM-001','관리자 대시보드','/admin','현재'],
    ['관리자','SCR-ADM-CUR-001','교과·교육과정','/admin/curriculum','현재/보완'],
    ['관리자','SCR-ADM-PRG-001','프로그램','/admin/programs','현재/보완'],
    ['관리자','SCR-ADM-CMP-001','이수판정','/admin/completion','현재/보완'],
    ['관리자','SCR-ADM-BEN-001','수혜','/admin/benefits','현재/보완'],
    ['관리자','SCR-ADM-COM-001','참여기업','/admin/companies','현재/보완'],
    ['관리자','SCR-ADM-BGT-001','예산','/admin/budget','현재/보완'],
    ['관리자','SCR-ADM-PER-001','성과','/admin/performance','현재/보완'],
    ['관리자','SCR-ADM-CMS-001','CMS','/admin/content','현재/보완'],
    ['관리자','SCR-ADM-FIL-001','파일관리','/admin/files','현재/보완'],
    ['관리자','SCR-ADM-AUD-001','감사로그','/admin/audit','현재/보완'],
    ['관리자','SCR-ADM-ROL-001','사용자·권한','/admin/access','목표'],
    ['관리자','SCR-ADM-QLT-001','품질·인수','/admin/quality','목표'],
    ['관리자','SCR-ADM-RSK-001','위험·잔여과제','/admin/risks','목표'],
    ['관리자','SCR-ADM-CHG-001','변경·승인','/admin/changes','목표'],
    ['관리자','SCR-ADM-OPS-001','운영·복구','/admin/operations','목표'],
    ['관리자','SCR-ADM-SET-001','시스템 설정','/admin/settings','현재/보완'],
]
appendices.append(Chapter('C', '전체 메뉴·화면·라우트', 'PARTIAL', [
    SectionBlock('라우트 목록', table=TableBlock(['포털','화면 ID','화면명','권장 경로','상태'], route_rows, [2.0,3.4,3.2,5.1,2.8])),
    SectionBlock('라우트 관리원칙', bullets=['최상위 화면과 상세화면은 동일 리소스 ID를 사용한다.', '메뉴 Registry와 Route Guard의 역할·Scope 설정을 동기화한다.', '현재 구현과 목표화면을 상태로 구분한다.', '감사 Deep Link는 삭제·보관된 리소스도 권한 범위에서 확인 가능해야 한다.']),
    SectionBlock('정합성 검토사항', bullets=['현재 라우트 수는 저장소 기준 재검산하여 기준선에 반영한다.', '/admin/tests와 /admin/quality 등 중복 제안경로는 /admin/quality로 통합한다.', '운영·위험·변경화면은 플랫폼 구현과 외부도구 연계를 구분한다.']),
]))

# D Traceability concise but comprehensive
trace_rows = [
    ['TRC-P0-001','대학 SSO','AUT-*','인증 Adapter·세션','실 IdP 로그인','BLOCKED_EXTERNAL'],
    ['TRC-P0-002','역할·Scope','ROL-*','Guard·Membership','IDOR·범위 E2E','PARTIAL'],
    ['TRC-P0-003','실 PostgreSQL','DAT-*','Migration·Repository','통합·동시성','REAL_ENV_REQUIRED'],
    ['TRC-P0-004','이수정책','CMP-*','규칙·Snapshot','업무 UAT','BLOCKED_EXTERNAL'],
    ['TRC-P0-005','장학금','BEN-*','정책·후보·지급','산식·대조','BLOCKED_EXTERNAL'],
    ['TRC-P0-006','예산 원장대조','BGT-*','ERP·RCMS','실원장 대조','BLOCKED_EXTERNAL'],
    ['TRC-P0-007','파일보안','FIL-*','S3·Scanner','감염·복구','REAL_ENV_REQUIRED'],
    ['TRC-P0-008','개인정보','PRV-*','동의·파기','개인정보 승인','BLOCKED_EXTERNAL'],
    ['TRC-P0-009','브라우저 E2E','TST-*','Playwright 등','역할별 E2E','NOT_IMPLEMENTED'],
    ['TRC-P0-010','백업·복구','OPS-*','DB·파일·Secret','전체 복구','REAL_ENV_REQUIRED'],
]
appendices.append(Chapter('D', '요구사항 추적표', 'PARTIAL', [
    SectionBlock('추적성 방향', code='사업계획·지침 → 요구사항 → 화면·API·데이터 → 구현 → 테스트 → 결함·변경 → 인수 → 운영증적'),
    SectionBlock('식별자 체계', table=TableBlock(['식별자','의미','예시'], [
        ['SRC-*','근거문서','SRC-PLAN-001'], ['REQ/영역-*','요구사항','BEN-PAY-001'], ['SCR-*','화면','SCR-ADM-BEN-001'], ['API-*','API','API-BEN-001'], ['DAT-*','데이터','DAT-BEN-001'], ['TST-*','테스트','TST-BEN-IT-001'], ['EVD-*','증적','EVD-BEN-001'], ['DEF-*','결함','DEF-SEC-001'], ['CHG-*','변경','CHG-CUR-001'], ['DEC-*','결정','DEC-AUT-001'], ['RSK-*','위험','RSK-FIL-001'], ['GAP-*','잔여과제','GAP-P0-001'],
    ], [3.4,6.0,7.1])),
    SectionBlock('P0 추적 요약', table=TableBlock(['추적 ID','핵심 요구사항','요구사항','구현대상','시험·증적','상태'], trace_rows, [2.7,3.2,2.5,3.5,3.3,2.3])),
    SectionBlock('완전성 기준', bullets=['모든 P0·P1은 출처·책임자·화면 또는 비화면 처리·API·데이터·테스트·인수기준을 가져야 한다.', '화면 없는 요구사항은 Job·설정·Runbook·모니터링에 연결한다.', '미실행·외부대기 항목을 커버리지에서 제외하지 않는다.', '구현상태와 인수상태를 분리한다.']),
]))

# E API/Data complete list
api_groups = [
    ['공통·시스템','8','Health, Session, Reference, System'], ['Academic','21','교과·개설·교육과정·Import'], ['Programs','17','프로그램·신청·운영·학습'], ['Completion','8','이수판정·경험·공개 포트폴리오'], ['Benefits','7','정책·후보·승인·지급'], ['Companies','20','기업·신청·담당자·활동·취업'], ['Budget','6','요약·변경·배정·집행'], ['Performance','12','지표·목표·결과·증빙·검토'], ['Content','6','공개·내부 콘텐츠·상태·버전'], ['Files','9','업로드·다운로드·보존·파기'], ['Audit','2','조회·CSV 내보내기'],
]
tables53 = ['business_years','terms','code_groups','code_values','users','roles','user_roles','students','course_masters','course_offerings','curricula','curriculum_requirements','import_jobs','import_staging_rows','audit_logs','file_retention_policies','stored_files','programs','program_sessions','program_applications','attendance_events','attendance_records','assignments','assignment_submissions','surveys','survey_responses','program_completions','course_completions','experiential_records','completion_assessments','benefit_policies','benefit_eligibility_rules','benefit_candidates','benefit_approvals','benefit_payments','company_applications','companies','company_contacts','company_experts','company_commitments','company_participations','budget_allocations','budget_executions','budget_change_history','performance_indicators','performance_targets','performance_results','performance_evidence','performance_reviews','content_items','content_attachments','content_versions','rate_limit_counters']
appendices.append(Chapter('E', 'API·데이터 목록', 'PARTIAL', [
    SectionBlock('API 현황', table=TableBlock(['영역','Endpoint 수','범위'], api_groups + [['합계','116','현재 기준 정적 집계']], [4.0,3.0,9.5])),
    SectionBlock('API 공통규칙', bullets=['업무 API는 /api/v1 경로를 사용한다.', '요청 ID와 공통 오류응답을 사용한다.', '보호 API는 역할·Scope·소유권을 서버에서 검증한다.', '금전·신청·제출 API는 멱등성을 적용한다.', '전체 Endpoint를 OpenAPI와 계약시험에 연결한다.']),
    SectionBlock('물리 테이블 53개', table=TableBlock(['순번','테이블'], [[str(i), t] for i,t in enumerate(tables53,1)], [2.0,14.5])),
    SectionBlock('주요 구조적 보완', bullets=['company_memberships 추가', 'user_roles PK에 복수 Scope 지원', '동의·권리요청·이의·환수·대조·제출 엔터티 추가', '범용 file_relationships 및 scan_results 추가', 'outbox·inbox·integration_operations 추가', 'release·backup·restore 운영 엔터티 추가']),
]))

# F checklist summary
f_rows = [
    ['정적·Build','Typecheck, Build, Secret Scan, SBOM, 이미지 취약점','P0','CI 증적'],
    ['API 계약','요청·응답·오류·페이지·Rate Limit','P0','계약시험'],
    ['인증·권한','SSO, 세션, 역할, Scope, IDOR, 자기승인','P0','실환경·E2E'],
    ['DB 통합','Migration, FK, Unique, Transaction, 동시성','P0','실 PostgreSQL'],
    ['업무기능','교과·프로그램·이수·수혜·기업·예산·성과','P0/P1','통합·UAT'],
    ['파일·보안','형식, Scanner, S3, 다운로드, 파기','P0','실환경'],
    ['비기능','성능, 접근성, 브라우저','P1','시험보고서'],
    ['운영','배포, 모니터링, Backup, Restore, RPO/RTO','P0','OAT·복구보고서'],
]
appendices.append(Chapter('F', '테스트·인수 체크리스트', 'PARTIAL / REAL_ENV_REQUIRED', [
    SectionBlock('테스트 유형', table=TableBlock(['영역','주요 확인사항','우선순위','필수 증적'], f_rows, [3.0,7.8,2.2,3.5])),
    SectionBlock('진입·종료기준', code='진입: 요구사항 승인 + Commit 고정 + 환경·데이터·외부정보 준비\n종료: P0 100% 통과 + Critical·Major 0건 + UAT·OAT 승인 + 전체 복구 통과'),
    SectionBlock('현재 판단', table=status_table('PARTIAL / REAL_ENV_REQUIRED', '실 PostgreSQL, 대학 SSO, 브라우저 E2E, ERP·RCMS, S3·Scanner, UAT·OAT', 'CI·통합·E2E·보안·성능·접근성·UAT·OAT·복구 보고서')),
    SectionBlock('Go/No-Go 기준', bullets=['P0 미실행·실패 1건이라도 있으면 NO-GO', 'Critical·Major 결함 0건', '개인정보·보안 승인 완료', 'Backup·Restore와 RPO·RTO 실측', '업무·운영 인수자 서명']),
]))

# G risks
risk_rows = [
    ['RSK-AUT-001','대학 SSO 미연계','5','5','25','BLOCKED_EXTERNAL'],
    ['RSK-ROL-001','역할 Scope·기업 Membership 불완전','4','5','20','PARTIAL'],
    ['RSK-PRV-001','개인정보 동의·처리체계 미완료','4','5','20','BLOCKED_EXTERNAL'],
    ['RSK-DAT-001','실 PostgreSQL 통합검증 미완료','4','5','20','REAL_ENV_REQUIRED'],
    ['RSK-BEN-001','장학금 정책·지급대조 미확정','4','5','20','BLOCKED_EXTERNAL'],
    ['RSK-BGT-001','ERP·RCMS 원장대조 미완료','4','5','20','BLOCKED_EXTERNAL'],
    ['RSK-FIL-001','S3·악성코드 검사 미검증','4','5','20','REAL_ENV_REQUIRED'],
    ['RSK-OPS-001','전체 복구 미검증','4','5','20','REAL_ENV_REQUIRED'],
    ['RSK-TST-001','브라우저 E2E 부재','4','4','16','NOT_IMPLEMENTED'],
    ['RSK-ORG-001','소규모 조직 역할집중','4','4','16','OPEN'],
]
p0_gaps = [
    ['GAP-P0-001','대학 SSO·중앙 세션'], ['GAP-P0-002','역할·Scope·기업 Membership'], ['GAP-P0-003','개인정보 처리목록·동의·권리·파기'], ['GAP-P0-004','실 PostgreSQL 통합 CI·동시성'], ['GAP-P0-005','공식 이수·장학금·성과정책'], ['GAP-P0-006','ERP·RCMS·전문기관 연계·대조'], ['GAP-P0-007','운영 S3·악성코드 검사·파일 Scope'], ['GAP-P0-008','브라우저 P0 E2E'], ['GAP-P0-009','Backup 자동화·전체 Restore'], ['GAP-P0-010','UAT·OAT·최종 Go 승인'],
]
appendices.append(Chapter('G', '위험·의사결정·잔여과제', 'PARTIAL / BLOCKED_EXTERNAL', [
    SectionBlock('위험평가', paragraphs=[p('발생 가능성 1~5와 영향도 1~5를 곱하여 위험점수를 산정한다. 20~25는 치명적, 15~19는 매우 높음으로 보아 원칙적으로 출시를 차단한다.')]),
    SectionBlock('핵심 위험', table=TableBlock(['위험 ID','위험','가능성','영향','점수','상태'], risk_rows, [2.7,5.6,1.7,1.7,1.7,3.1])),
    SectionBlock('P0 잔여과제', table=TableBlock(['Gap ID','과제'], p0_gaps, [4.0,12.5])),
    SectionBlock('주요 의사결정', table=TableBlock(['결정 ID','안건','현재 상태'], [
        ['DEC-AUT-001','대학 인증 프로토콜·Claim','BLOCKED_EXTERNAL'], ['DEC-ROL-001','역할·Scope·직무분리','PENDING_APPROVAL'], ['DEC-BEN-001','장학금 자격·산식·환수','PENDING_APPROVAL'], ['DEC-BGT-001','ERP·RCMS 연계방식','BLOCKED_EXTERNAL'], ['DEC-PER-001','성과 산식·제출·인정값','BLOCKED_EXTERNAL'], ['DEC-PRV-001','개인정보 수집·보존·제공','PENDING_APPROVAL'], ['DEC-OPS-001','Hosting·S3·Scanner·로그','PENDING_APPROVAL'], ['DEC-OPS-002','RPO·RTO·DR','PENDING_APPROVAL'],
    ], [3.2,8.5,4.8])),
    SectionBlock('현재 판정', table=status_table('치명적 / NO-GO', 'P0 다수, 외부정보·실환경·인수 미완료', 'P0 0건, 최종 위험평가, 위험수용·Go 승인')),
]))

# H changes
appendices.append(Chapter('H', '변경·승인기록', 'PARTIAL', [
    SectionBlock('변경 흐름', code='변경요청 → 영향분석 → 위험평가 → 승인 → 구현 → 시험 → 배포 → 확인 → 문서·기준선 갱신'),
    SectionBlock('변경등급', table=TableBlock(['등급','기준','예시'], [
        ['L0','업무·데이터·보안 영향 없음','오탈자'], ['L1','제한적 UI·편의기능','정렬·문구'], ['L2','단일 기능·API·데이터 영향','신청항목'], ['L3','다중 도메인·개인정보·금액','장학정책'], ['L4','인증·대량 PII·지급·예산·전체 DB','SSO·Migration'],
    ], [2.2,7.0,7.3])),
    SectionBlock('현재 기준선·정정', table=TableBlock(['변경 ID','변경내용','상태'], [
        ['CHG-DOC-001','물리 테이블 총계 52개에서 53개로 정정(program_completions 포함)','DOCUMENT_BASELINE'], ['CHG-DOC-002','API 총계 115개에서 116개로 정정(benefit-payments 포함)','DOCUMENT_BASELINE'], ['CHG-CI-001','Windows CI 시험 실행 수정 커밋 2346fb2...','IMPLEMENTED_BASELINE'], ['CHG-REL-001','최초 운영 Go-Live','NO-GO'],
    ], [3.4,9.3,3.8])),
    SectionBlock('승인단계', bullets=['A1 업무·요구사항', 'A2 개인정보·보안', 'A3 기술설계', 'A4 구현·시험', 'A5 배포', 'A6 UAT', 'A7 OAT', 'A8 최종 Go']),
    SectionBlock('긴급변경', bullets=['서비스중단·개인정보·보안·금전·데이터 위험에 한해 최소범위로 수행한다.', '즉시 검증 후 사후 코드검토·회귀시험·승인을 완료한다.', '일정관리 실패나 일반 개선은 긴급변경 사유가 아니다.']),
    SectionBlock('문서 상태', paragraphs=[p('현재 통합 명세서는 v0.9.0 검토본이며 공식 문서 소유자·검토자·승인자·승인일이 확정되지 않았다. 최종 승인 후 v1.0.0 기준선으로 전환하고 릴리스 패키지와 연결한다.')]),
]))

# ---------------------------------------------------------------------------
# Extra registers for completeness
# ---------------------------------------------------------------------------
external_requests = [
    ['EXT-SSO-001','대학 정보전산원','Issuer, Client, Redirect/Logout URI, Claim, 시험계정, MFA·세션정책','P0'],
    ['EXT-SIS-001','학사부서·정보전산원','학생·학적·교과·개설·수강·성적·정정 Interface','P0'],
    ['EXT-ERP-001','장학·재무부서','지급요청·결과·취소·환수·중복대조','P0'],
    ['EXT-RCMS-001','회계·산학협력단','집행·정산·불인정·환수·참조번호','P0'],
    ['EXT-KIAT-001','전문기관','성과 정의·제출 Schema·접수·보완·인정값','P0'],
    ['EXT-PRV-001','개인정보 담당','처리목록·동의문·보존·제공·사고 연락망','P0'],
    ['EXT-OPS-001','운영·보안부서','Hosting, DNS/TLS, Secret, S3, Scanner, 로그·Metrics, Backup·DR','P0'],
]

# ---------------------------------------------------------------------------
# Markdown writer
# ---------------------------------------------------------------------------
def md_table(tb: TableBlock) -> str:
    out = ['| ' + ' | '.join(tb.headers) + ' |', '| ' + ' | '.join(['---']*len(tb.headers)) + ' |']
    for row in tb.rows:
        out.append('| ' + ' | '.join(str(x).replace('|','\\|').replace('\n','<br>') for x in row) + ' |')
    if tb.note:
        out.append('')
        out.append(f'> {tb.note}')
    return '\n'.join(out)


def chapter_to_md(ch: Chapter, appendix=False) -> str:
    head = f'# 부록 {ch.number}. {ch.title}' if appendix else f'# 제{ch.number}장. {ch.title}'
    lines=[head, '', f'**작성상태:** `{ch.status}`', '']
    for idx,s in enumerate(ch.sections,1):
        lines.append(f'## {ch.number}.{idx} {s.title}' if not appendix else f'## {ch.number}.{idx} {s.title}')
        lines.append('')
        for para in s.paragraphs:
            lines += [para,'']
        for b in s.bullets:
            lines.append(f'- {b}')
        if s.bullets: lines.append('')
        if s.code:
            lines += ['```text', s.code, '```', '']
        if s.table:
            lines += [md_table(s.table),'']
    return '\n'.join(lines)

md_lines = [
    '# 국립한국교통대학교 첨단산업 인재양성 부트캠프사업(AI) 통합운영 플랫폼 세부 명세서',
    '',
    '- 버전: v0.9.0 검토본',
    '- 작성기준일: 2026-07-30',
    '- 대상 저장소: `Nawii06/bootcampAI_page`',
    '- 기준 커밋: `2346fb2ed634c94f6d7d978dbda1ec8fe22d93ef` (저장소 반영 전 재확인)',
    '- 현재 출시판정: **NO-GO**',
    '',
    '> 본 문서는 공식 승인 전 통합 검토본이다. 정책·외부연계·실환경 검증사항은 확정상태와 구분한다.',
    '',
    '## 목차',
    '',
]
for ch in chapters:
    md_lines.append(f'- 제{ch.number}장 {ch.title}')
for ap in appendices:
    md_lines.append(f'- 부록 {ap.number} {ap.title}')
md_lines += ['', '---', '']
for ch in chapters:
    md_lines += [chapter_to_md(ch), '', '---', '']
for ap in appendices:
    md_lines += [chapter_to_md(ap, True), '', '---', '']
md_lines += ['# 외부정보 요청 요약', '', md_table(TableBlock(['외부 ID','요청대상','요청사항','우선순위'], external_requests)), '', '# 최종 판정', '', '현재 문서 구성은 완료되었으나 대학 SSO, 역할·Scope, 개인정보, 공식 정책, 실 PostgreSQL, 외부 원장, 파일 보안, E2E, UAT·OAT 및 전체 복구 증적이 미완료이므로 운영 출시판정은 `NO-GO`이다.', '']
MD_PATH.write_text('\n'.join(md_lines), encoding='utf-8')

# ---------------------------------------------------------------------------
# DOCX helpers
# ---------------------------------------------------------------------------
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is None:
        shd = OxmlElement('w:shd')
        tcPr.append(shd)
    shd.set(qn('w:fill'), fill)


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top',top),('start',start),('bottom',bottom),('end',end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v)); node.set(qn('w:type'),'dxa')


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    tblHeader.set(qn('w:val'), 'true')
    trPr.append(tblHeader)


def set_table_borders(table, color=MEDGRAY, size='4'):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in('w:tblBorders')
    if borders is None:
        borders = OxmlElement('w:tblBorders')
        tblPr.append(borders)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = f'w:{edge}'
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), size)
        element.set(qn('w:space'), '0')
        element.set(qn('w:color'), color)


def set_run_font(run, name=FONT_BODY, size=9.4, bold=False, color=CHARCOAL):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para_font(paragraph, name=FONT_BODY, size=9.4, bold=False, color=CHARCOAL):
    for run in paragraph.runs:
        set_run_font(run, name, size, bold, color)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run('페이지 ')
    set_run_font(run, FONT_BODY, 8, False, DARKGRAY)
    fldChar1 = OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText'); instrText.set(qn('xml:space'), 'preserve'); instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1); run._r.append(instrText); run._r.append(fldChar2)


def add_table(doc: Document, tb: TableBlock):
    table = doc.add_table(rows=1, cols=len(tb.headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for j,h in enumerate(tb.headers):
        cell=hdr.cells[j]
        set_cell_shading(cell, BLUE)
        set_cell_margins(cell)
        para=cell.paragraphs[0]
        para.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=para.add_run(str(h))
        set_run_font(run, FONT_HEAD, 8.6, True, WHITE)
        cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for r_idx,row in enumerate(tb.rows):
        cells=table.add_row().cells
        for j in range(len(tb.headers)):
            txt = str(row[j]) if j < len(row) else ''
            cell=cells[j]
            set_cell_margins(cell)
            if r_idx % 2 == 1:
                set_cell_shading(cell, SOFTWHITE)
            para=cell.paragraphs[0]
            para.paragraph_format.space_after=Pt(0)
            para.paragraph_format.line_spacing=1.0
            run=para.add_run(txt)
            set_run_font(run, FONT_BODY, 8.15, False, CHARCOAL)
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if tb.widths:
        for row in table.rows:
            for j,w in enumerate(tb.widths):
                row.cells[j].width = Cm(w)
    # Keep small tables together to avoid one-row spill pages. Large inventories may paginate normally.
    if len(tb.rows) <= 8:
        for row in table.rows[:-1]:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.keep_with_next = True
    if tb.note:
        para=doc.add_paragraph()
        para.paragraph_format.space_before=Pt(2)
        para.paragraph_format.space_after=Pt(5)
        run=para.add_run('주: ' + tb.note)
        set_run_font(run, FONT_BODY, 8, False, DARKGRAY)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)
    return table


def add_callout(doc: Document, text: str, accent=BLUE):
    table=doc.add_table(rows=1, cols=1)
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color=accent, size='6')
    cell=table.cell(0,0); set_cell_shading(cell, SOFTWHITE); set_cell_margins(cell,120,160,120,160)
    para=cell.paragraphs[0]
    run=para.add_run(text)
    set_run_font(run, FONT_BODY, 9, True, accent)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)


def add_code(doc: Document, text: str):
    table=doc.add_table(rows=1, cols=1)
    set_table_borders(table, color=MEDGRAY, size='4')
    cell=table.cell(0,0); set_cell_shading(cell, SOFTWHITE); set_cell_margins(cell,100,140,100,140)
    para=cell.paragraphs[0]
    for i,line in enumerate(text.split('\n')):
        if i: para.add_run().add_break()
        run=para.add_run(line)
        set_run_font(run, 'NanumGothicCoding', 8.5, False, CHARCOAL)
    doc.add_paragraph().paragraph_format.space_after=Pt(1)


def add_section_block(doc: Document, num: str, idx: int, s: SectionBlock, appendix=False):
    heading = doc.add_paragraph(style='Heading 2')
    heading.paragraph_format.keep_with_next=True
    if s.title == '요구사항 목록' and str(num).isdigit() and int(num) >= 5:
        heading.paragraph_format.page_break_before = True
    heading.paragraph_format.space_before=Pt(10)
    heading.paragraph_format.space_after=Pt(4)
    run=heading.add_run(f'{num}.{idx} {s.title}')
    set_run_font(run, FONT_HEAD, 12.2, True, BLUE)
    for para_text in s.paragraphs:
        para=doc.add_paragraph(style='Body Text')
        para.paragraph_format.space_after=Pt(4)
        para.paragraph_format.line_spacing=1.18
        para.paragraph_format.first_line_indent=Cm(0.65)
        run=para.add_run(para_text)
        set_run_font(run, FONT_BODY, 9.3, False, CHARCOAL)
    for b in s.bullets:
        para=doc.add_paragraph(style='List Bullet')
        para.paragraph_format.left_indent=Cm(0.55)
        para.paragraph_format.first_line_indent=Cm(-0.25)
        para.paragraph_format.space_after=Pt(2)
        run=para.add_run(b)
        set_run_font(run, FONT_BODY, 9.1, False, CHARCOAL)
    if s.code:
        add_code(doc,s.code)
    if s.table:
        add_table(doc,s.table)


def add_chapter(doc: Document, ch: Chapter, appendix=False):
    head=doc.add_paragraph()
    head.paragraph_format.page_break_before = True
    head.paragraph_format.space_after=Pt(8)
    run=head.add_run((f'부록 {ch.number}. ' if appendix else f'제{ch.number}장. ') + ch.title)
    set_run_font(run, FONT_HEAD, 20, True, BLUE)
    # status line
    para=doc.add_paragraph()
    para.paragraph_format.space_after=Pt(10)
    r=para.add_run('작성상태  '); set_run_font(r,FONT_HEAD,8.5,True,DARKGRAY)
    r=para.add_run(ch.status); set_run_font(r,FONT_BODY,8.5,True,BURGUNDY if 'NO-GO' in ch.status or 'BLOCKED' in ch.status else BLUE)
    for idx,s in enumerate(ch.sections,1):
        add_section_block(doc,ch.number,idx,s,appendix)


doc=Document()
sec=doc.sections[0]
sec.page_width=Cm(21.0); sec.page_height=Cm(29.7)
sec.top_margin=Cm(1.8); sec.bottom_margin=Cm(1.6); sec.left_margin=Cm(1.8); sec.right_margin=Cm(1.8)
sec.header_distance=Cm(0.7); sec.footer_distance=Cm(0.7)

# styles
styles=doc.styles
styles['Normal'].font.name=FONT_BODY
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),FONT_BODY)
styles['Normal'].font.size=Pt(9.4)
for st_name, size, color in [('Title',26,BLUE),('Heading 1',18,BLUE),('Heading 2',12.2,BLUE),('Heading 3',10.5,BURGUNDY)]:
    st=styles[st_name]
    st.font.name=FONT_HEAD; st._element.rPr.rFonts.set(qn('w:eastAsia'),FONT_HEAD); st.font.size=Pt(size); st.font.color.rgb=RGBColor.from_string(color); st.font.bold=True
styles['Body Text'].font.name=FONT_BODY; styles['Body Text']._element.rPr.rFonts.set(qn('w:eastAsia'),FONT_BODY); styles['Body Text'].font.size=Pt(9.3)
styles['List Bullet'].font.name=FONT_BODY; styles['List Bullet']._element.rPr.rFonts.set(qn('w:eastAsia'),FONT_BODY); styles['List Bullet'].font.size=Pt(9.1)

# header/footer
header=sec.header
hp=header.paragraphs[0]
hp.alignment=WD_ALIGN_PARAGRAPH.LEFT
r=hp.add_run('부트캠프(AI) 통합운영 플랫폼 세부 명세서  |  v0.9.0 검토본')
set_run_font(r,FONT_HEAD,8,True,BLUE)
add_page_number(sec.footer.paragraphs[0])

# Cover
cover=doc.add_paragraph()
cover.alignment=WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before=Pt(90)
r=cover.add_run('국립한국교통대학교')
set_run_font(r,FONT_HEAD,20,True,BLUE)
r.add_break()
r=cover.add_run('첨단산업 인재양성 부트캠프사업(AI)')
set_run_font(r,FONT_HEAD,18,True,BLUE)
r.add_break()
r=cover.add_run('통합운영 플랫폼 세부 명세서')
set_run_font(r,FONT_HEAD,28,True,BLUE)
sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER; sub.paragraph_format.space_before=Pt(18)
r=sub.add_run('Integrated Operations Platform Detailed Specification')
set_run_font(r,FONT_BODY,11,False,DARKGRAY)
add_callout(doc,'v0.9.0 통합 검토본 · 공식 승인 전 · 현재 출시판정 NO-GO',BURGUNDY)
meta=doc.add_paragraph(); meta.alignment=WD_ALIGN_PARAGRAPH.CENTER; meta.paragraph_format.space_before=Pt(120)
for line in ['작성기준일  2026. 7. 30.','대상 저장소  Nawii06/bootcampAI_page','기준 커밋  2346fb2ed634c94f6d7d978dbda1ec8fe22d93ef','부트캠프(AI)사업단']:
    r=meta.add_run(line); set_run_font(r,FONT_BODY,10,False,CHARCOAL); r.add_break()

# document control
doc.add_page_break()
h=doc.add_paragraph(); r=h.add_run('문서 통제정보'); set_run_font(r,FONT_HEAD,20,True,BLUE)
add_table(doc, TableBlock(['항목','내용'], [
    ['문서 소유자','확정 필요'], ['작성자','부트캠프(AI)사업단 실무 검토본'], ['검토자','확정 필요'], ['승인자','확정 필요'], ['보안등급','내부업무용(최종 확정 필요)'], ['문서 버전','v0.9.0'], ['문서 상태','REVIEW / 공식 승인 전'], ['운영 상태','PARTIAL / BLOCKED_EXTERNAL / REAL_ENV_REQUIRED'], ['출시판정','NO-GO'],
], [4.0,12.5]))
h=doc.add_paragraph(); r=h.add_run('변경이력'); set_run_font(r,FONT_HEAD,15,True,BLUE)
add_table(doc, TableBlock(['버전','변경일','변경내용','상태'], [
    ['0.1.0','2026-07','초기 홈페이지 통합관리 변경명세','DRAFT'], ['0.9.0','2026-07-30','본문 25장 및 부록 A-H 통합·수치정정·기준선 정리','REVIEW'], ['1.0.0','미정','공식 검토·승인 및 외부정보 반영','예정'],
], [2.5,3.0,8.0,3.0]))
h=doc.add_paragraph(); r=h.add_run('승인란'); set_run_font(r,FONT_HEAD,15,True,BLUE)
add_table(doc, TableBlock(['구분','성명·직위','검토·승인 의견','일자·서명'], [['작성','','',''],['업무 검토','','',''],['개인정보·보안 검토','','',''],['운영 검토','','',''],['최종 승인','','','']], [2.4,4.0,6.8,3.3]))

# TOC static
doc.add_page_break()
h=doc.add_paragraph(); r=h.add_run('목차'); set_run_font(r,FONT_HEAD,22,True,BLUE)
for ch in chapters:
    para=doc.add_paragraph(); para.paragraph_format.space_after=Pt(2); r=para.add_run(f'제{ch.number}장  {ch.title}'); set_run_font(r,FONT_BODY,9.4,False,CHARCOAL)
for ap in appendices:
    para=doc.add_paragraph(); para.paragraph_format.space_after=Pt(2); r=para.add_run(f'부록 {ap.number}  {ap.title}'); set_run_font(r,FONT_BODY,9.4,False,CHARCOAL)

# Executive summary
doc.add_page_break()
h=doc.add_paragraph(); r=h.add_run('요약 및 현재 판단'); set_run_font(r,FONT_HEAD,22,True,BLUE)
for text in [
    '본 명세서는 사업 홍보 홈페이지를 학생·기업·관리자 업무와 교과·프로그램·이수·수혜·기업·예산·성과·CMS·파일·감사 기능을 연결하는 통합운영 플랫폼으로 전환하기 위한 기준을 정의한다.',
    '현재 저장소는 116개 API Endpoint와 53개 물리테이블을 중심으로 상당한 기능기반을 갖추고 있으나, 실제 대학 SSO, 역할 Scope와 기업 Membership, 개인정보 운영정책, 학사·ERP·RCMS·전문기관 연계, 실 PostgreSQL·S3·악성코드 검사, 브라우저 E2E, UAT·OAT 및 전체 복구 증적이 부족하다.',
    '따라서 현재 정식운영 판정은 NO-GO이며, 모든 P0 Gap을 증적으로 종료한 후에만 Go 검토를 수행한다.'
]:
    para=doc.add_paragraph(); para.paragraph_format.space_after=Pt(5); para.paragraph_format.line_spacing=1.2; r=para.add_run(text); set_run_font(r,FONT_BODY,10,False,CHARCOAL)
add_table(doc, TableBlock(['구분','현재 기준'], [
    ['본문·부록','제1장~제25장, 부록 A~H'], ['현재 API','116개 Endpoint'], ['현재 DB','53개 물리테이블'], ['코드 기준선','2346fb2ed634c94f6d7d978dbda1ec8fe22d93ef'], ['핵심 P0','SSO, Scope, 개인정보, 정책, DB, 외부원장, 파일보안, E2E, 복구, 인수'], ['출시판정','NO-GO'],
], [4.2,12.3]))

for ch in chapters:
    add_chapter(doc,ch,False)
for ap in appendices:
    add_chapter(doc,ap,True)

# External request and final checklist
doc.add_page_break()
h=doc.add_paragraph(); r=h.add_run('외부정보 요청 종합표'); set_run_font(r,FONT_HEAD,20,True,BLUE)
add_table(doc, TableBlock(['외부 ID','요청대상','필요정보','우선순위'], external_requests, [3.2,4.0,7.2,2.1]))
h=doc.add_paragraph(); r=h.add_run('운영 전 최종 확인'); set_run_font(r,FONT_HEAD,15,True,BLUE)
for item in ['대학 SSO 실연계 및 세션 회수', '역할·Scope·기업 Membership 및 IDOR 전수시험', '개인정보 처리목록·동의·보존·파기 승인', '공식 이수·장학금·예산·성과 정책 승인', '실 PostgreSQL Migration·통합·동시성 시험', 'ERP·RCMS·전문기관 Sandbox·대조', '운영 S3·악성코드 검사·파일 복구', '브라우저 P0 E2E·보안·성능·접근성', 'Backup 자동화와 빈 환경 전체 Restore', 'UAT·OAT·최종 위험평가·Go 승인']:
    para=doc.add_paragraph(style='List Bullet'); r=para.add_run(item); set_run_font(r,FONT_BODY,9.4,False,CHARCOAL)
add_callout(doc,'최종 판정: 현재 NO-GO. P0 잔여과제와 실환경·인수 증적 완료 후 재판정한다.',BURGUNDY)

# properties
props=doc.core_properties
props.title='국립한국교통대학교 부트캠프(AI) 통합운영 플랫폼 세부 명세서'
props.subject='본문 제1장~제25장 및 부록 A~H 통합 명세'
props.author='국립한국교통대학교 부트캠프(AI)사업단'
props.keywords='부트캠프, AI, 통합운영 플랫폼, 세부 명세서, 요구사항, API, 데이터, 테스트, 운영'
props.comments='v0.9.0 검토본. 공식 승인 전.'

doc.save(DOCX_PATH)

manifest={
    'document': DOCX_PATH.name,
    'markdown_source': MD_PATH.name,
    'version':'0.9.0',
    'date':'2026-07-30',
    'repository':'Nawii06/bootcampAI_page',
    'baseline_commit':'2346fb2ed634c94f6d7d978dbda1ec8fe22d93ef',
    'chapters': [{'number':c.number,'title':c.title,'status':c.status} for c in chapters],
    'appendices': [{'number':c.number,'title':c.title,'status':c.status} for c in appendices],
    'api_endpoint_count':116,
    'physical_table_count':53,
    'release_decision':'NO-GO',
}
MANIFEST_PATH.write_text(json.dumps(manifest,ensure_ascii=False,indent=2),encoding='utf-8')
print(DOCX_PATH)
print(MD_PATH)
print(MANIFEST_PATH)
